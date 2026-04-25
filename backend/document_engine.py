from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass, field
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


COVER_LETTER_MAX_PARAGRAPHS = 4
COVER_LETTER_TARGET_TOTAL_CHARS = 1100
MAX_PLAN_REPAIR_ATTEMPTS = 5
SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[.!?])\s+")


@dataclass(frozen=True)
class ThemeSpec:
    theme_id: str
    density: str
    ats_profile: str
    body_font: str
    body_size_pt: float
    body_color: tuple[int, int, int]
    heading_font: str
    heading_size_pt: float
    heading_color: tuple[int, int, int]
    name_size_pt: float
    title_size_pt: float
    title_color: tuple[int, int, int]
    top_margin_in: float
    right_margin_in: float
    bottom_margin_in: float
    left_margin_in: float
    line_spacing: float
    paragraph_after_pt: float
    section_before_pt: float
    section_after_pt: float
    bullet_indent_in: float
    max_resume_experiences: int
    max_bullets_per_experience: int
    summary_target_chars: int
    skills_target_chars: int
    resume_header_alignment: str = "center"
    cover_letter_date_alignment: str = "left"
    heading_case: str = "title"
    title_italic: bool = False
    header_divider: bool = False
    divider_color: tuple[int, int, int] = (200, 200, 200)


@dataclass
class DocumentPlan:
    doc_type: str
    page_budget: int
    theme_id: str
    density: str
    normalized_sections: dict[str, Any]
    section_order: list[str]
    layout_metrics: dict[str, Any]
    verification: dict[str, Any]
    repair_history: list[dict[str, Any]] = field(default_factory=list)
    attempt_count: int = 1


@dataclass
class VerificationIssue:
    code: str
    message: str
    actual: Any | None = None
    limit: Any | None = None


@dataclass
class RepairAction:
    action: str
    reason: str
    details: dict[str, Any] = field(default_factory=dict)


THEMES: dict[str, ThemeSpec] = {
    "classic_professional": ThemeSpec(
        theme_id="classic_professional",
        density="balanced",
        ats_profile="safe",
        body_font="Calibri",
        body_size_pt=11.0,
        body_color=(0, 0, 0),
        heading_font="Calibri",
        heading_size_pt=13.0,
        heading_color=(79, 129, 189),
        name_size_pt=23.0,
        title_size_pt=13.0,
        title_color=(90, 90, 90),
        top_margin_in=0.8,
        right_margin_in=1.0,
        bottom_margin_in=0.8,
        left_margin_in=1.0,
        line_spacing=1.02,
        paragraph_after_pt=4.0,
        section_before_pt=10.0,
        section_after_pt=4.0,
        bullet_indent_in=0.23,
        max_resume_experiences=3,
        max_bullets_per_experience=2,
        summary_target_chars=320,
        skills_target_chars=220,
        resume_header_alignment="center",
        cover_letter_date_alignment="left",
        heading_case="title",
        title_italic=False,
    ),
    "technical_compact": ThemeSpec(
        theme_id="technical_compact",
        density="compact",
        ats_profile="safe",
        body_font="Calibri",
        body_size_pt=10.5,
        body_color=(15, 15, 15),
        heading_font="Calibri",
        heading_size_pt=12.0,
        heading_color=(58, 110, 165),
        name_size_pt=22.0,
        title_size_pt=12.0,
        title_color=(85, 85, 85),
        top_margin_in=0.65,
        right_margin_in=0.9,
        bottom_margin_in=0.65,
        left_margin_in=0.9,
        line_spacing=1.0,
        paragraph_after_pt=2.0,
        section_before_pt=8.0,
        section_after_pt=3.0,
        bullet_indent_in=0.18,
        max_resume_experiences=4,
        max_bullets_per_experience=2,
        summary_target_chars=260,
        skills_target_chars=260,
        resume_header_alignment="center",
        cover_letter_date_alignment="left",
        heading_case="title",
        title_italic=False,
    ),
    "executive_clean": ThemeSpec(
        theme_id="executive_clean",
        density="balanced",
        ats_profile="safe",
        body_font="Cambria",
        body_size_pt=10.8,
        body_color=(26, 31, 36),
        heading_font="Cambria",
        heading_size_pt=11.0,
        heading_color=(52, 73, 94),
        name_size_pt=24.0,
        title_size_pt=11.5,
        title_color=(88, 95, 104),
        top_margin_in=0.85,
        right_margin_in=1.0,
        bottom_margin_in=0.8,
        left_margin_in=1.0,
        line_spacing=1.04,
        paragraph_after_pt=4.0,
        section_before_pt=12.0,
        section_after_pt=4.0,
        bullet_indent_in=0.22,
        max_resume_experiences=3,
        max_bullets_per_experience=2,
        summary_target_chars=290,
        skills_target_chars=200,
        resume_header_alignment="left",
        cover_letter_date_alignment="right",
        heading_case="upper",
        title_italic=True,
    ),
    "ats_minimal": ThemeSpec(
        theme_id="ats_minimal",
        density="balanced",
        ats_profile="strict",
        body_font="Arial",
        body_size_pt=11.0,
        body_color=(0, 0, 0),
        heading_font="Arial",
        heading_size_pt=11.0,
        heading_color=(0, 0, 0),
        name_size_pt=20.0,
        title_size_pt=11.0,
        title_color=(32, 32, 32),
        top_margin_in=0.8,
        right_margin_in=0.95,
        bottom_margin_in=0.8,
        left_margin_in=0.95,
        line_spacing=1.0,
        paragraph_after_pt=3.0,
        section_before_pt=8.0,
        section_after_pt=3.0,
        bullet_indent_in=0.2,
        max_resume_experiences=3,
        max_bullets_per_experience=2,
        summary_target_chars=260,
        skills_target_chars=210,
        resume_header_alignment="left",
        cover_letter_date_alignment="left",
        heading_case="title",
        title_italic=False,
    ),
    "modern_minimal": ThemeSpec(
        theme_id="modern_minimal",
        density="balanced",
        ats_profile="safe",
        body_font="Calibri",
        body_size_pt=10.8,
        body_color=(24, 30, 36),
        heading_font="Calibri",
        heading_size_pt=11.5,
        heading_color=(36, 92, 128),
        name_size_pt=22.5,
        title_size_pt=11.0,
        title_color=(84, 96, 108),
        top_margin_in=0.82,
        right_margin_in=0.95,
        bottom_margin_in=0.78,
        left_margin_in=0.95,
        line_spacing=1.03,
        paragraph_after_pt=3.0,
        section_before_pt=11.0,
        section_after_pt=4.0,
        bullet_indent_in=0.2,
        max_resume_experiences=3,
        max_bullets_per_experience=2,
        summary_target_chars=280,
        skills_target_chars=220,
        resume_header_alignment="left",
        cover_letter_date_alignment="right",
        heading_case="title",
        title_italic=False,
        header_divider=True,
        divider_color=(36, 92, 128),
    ),
}


def _clean_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _contains_phrase(text: str, phrase: str) -> bool:
    pattern = r"\b" + r"\s+".join(re.escape(part) for part in phrase.split()) + r"\b"
    return re.search(pattern, text) is not None


def _list_to_text(value) -> str:
    if isinstance(value, list):
        return ", ".join(str(item).strip() for item in value if str(item).strip())
    if value in (None, ""):
        return ""
    return str(value).strip()


def _trim_text(text: str, max_chars: int) -> str:
    cleaned = _clean_whitespace(text)
    if len(cleaned) <= max_chars:
        return cleaned

    shortened = cleaned[:max_chars].rstrip()
    sentence_end = max(shortened.rfind(". "), shortened.rfind("! "), shortened.rfind("? "))
    if sentence_end > max_chars // 2:
        return shortened[:sentence_end + 1].strip()

    shortened = shortened.rsplit(" ", 1)[0].rstrip(",;: ")
    if shortened and shortened[-1] not in ".!?":
        shortened += "."
    return shortened


def _compact_cover_letter_paragraph(text: str, *, max_sentences: int, max_chars: int) -> str:
    cleaned = _clean_whitespace(text)
    if not cleaned:
        return ""

    sentences = SENTENCE_BOUNDARY_RE.split(cleaned)
    if len(sentences) > max_sentences:
        cleaned = " ".join(sentences[:max_sentences]).strip()

    return _trim_text(cleaned, max_chars)


def _normalize_cover_letter_paragraphs(paragraphs) -> list[str]:
    if isinstance(paragraphs, str):
        items = [paragraphs]
    elif isinstance(paragraphs, list):
        items = [str(paragraph) for paragraph in paragraphs]
    else:
        paragraph_text = _list_to_text(paragraphs)
        items = [paragraph_text] if paragraph_text else []

    compacted = [
        _clean_whitespace(paragraph)
        for paragraph in items
        if _clean_whitespace(paragraph)
    ][:COVER_LETTER_MAX_PARAGRAPHS]

    if not compacted:
        return []

    normalized = []
    for index, paragraph in enumerate(compacted):
        is_closing_paragraph = index == len(compacted) - 1 and any(
            token in paragraph.lower()
            for token in ("thank", "appreciate", "welcome the opportunity", "look forward")
        )
        normalized.append(
            _compact_cover_letter_paragraph(
                paragraph,
                max_sentences=2 if is_closing_paragraph else 3,
                max_chars=220 if is_closing_paragraph else 360,
            )
        )

    total_chars = sum(len(paragraph) for paragraph in normalized)
    while total_chars > COVER_LETTER_TARGET_TOTAL_CHARS:
        longest_index = max(range(len(normalized)), key=lambda idx: len(normalized[idx]))
        current = normalized[longest_index]
        tighter_limit = max(180, len(current) - 80)
        updated = _compact_cover_letter_paragraph(
            current,
            max_sentences=2,
            max_chars=tighter_limit,
        )
        if updated == current:
            break
        normalized[longest_index] = updated
        total_chars = sum(len(paragraph) for paragraph in normalized)

    return [paragraph for paragraph in normalized if paragraph]


def _has_leadership_signal(normalized_sections: dict) -> bool:
    leadership_markers = (
        "staff",
        "principal",
        "lead",
        "director",
        "manager",
        "head of",
        "leadership",
        "mentored",
        "mentoring",
        "strategy",
    )
    text_parts = [
        normalized_sections.get("title", ""),
        normalized_sections.get("summary", ""),
        normalized_sections.get("role", ""),
    ]
    for experience in normalized_sections.get("experiences", []):
        if not isinstance(experience, dict):
            continue
        text_parts.append(str(experience.get("role", "")))
        text_parts.extend(str(bullet) for bullet in experience.get("bullets", []))
    combined = " ".join(_clean_whitespace(part).lower() for part in text_parts if part)
    return any(_contains_phrase(combined, marker) for marker in leadership_markers)


def _has_ats_simplicity_signal(normalized_sections: dict) -> bool:
    ats_markers = (
        "analyst",
        "accountant",
        "coordinator",
        "administrator",
        "specialist",
        "operations",
        "compliance",
        "government",
        "public sector",
        "healthcare",
        "support",
        "customer service",
        "case manager",
    )
    text_parts = [
        normalized_sections.get("title", ""),
        normalized_sections.get("summary", ""),
        normalized_sections.get("role", ""),
        normalized_sections.get("company", ""),
    ]
    combined = " ".join(_clean_whitespace(part).lower() for part in text_parts if part)
    return any(_contains_phrase(combined, marker) for marker in ats_markers)


def _has_design_signal(normalized_sections: dict) -> bool:
    design_markers = (
        "designer",
        "design",
        "ux",
        "ui",
        "product design",
        "brand",
        "visual",
        "creative",
        "art director",
        "interaction",
        "service design",
        "motion",
    )
    text_parts = [
        normalized_sections.get("title", ""),
        normalized_sections.get("summary", ""),
        normalized_sections.get("role", ""),
        normalized_sections.get("company", ""),
    ]
    for experience in normalized_sections.get("experiences", []):
        if not isinstance(experience, dict):
            continue
        text_parts.append(str(experience.get("role", "")))
        text_parts.extend(str(bullet) for bullet in experience.get("bullets", []))
    combined = " ".join(_clean_whitespace(part).lower() for part in text_parts if part)
    return any(_contains_phrase(combined, marker) for marker in design_markers)


def _requested_layout_strategy(normalized_sections: dict) -> str | None:
    strategy = str(normalized_sections.get("layout_strategy", "")).strip().lower()
    if strategy in {"ats_safe", "balanced", "executive", "compact", "creative_safe"}:
        return strategy
    return None


def _fit_cover_letter_paragraphs(
    paragraphs,
    *,
    paragraph_cap: int = COVER_LETTER_MAX_PARAGRAPHS,
    total_target_chars: int = COVER_LETTER_TARGET_TOTAL_CHARS,
    body_max_sentences: int = 3,
    closing_max_sentences: int = 2,
    body_max_chars: int = 360,
    closing_max_chars: int = 220,
) -> list[str]:
    compacted = _normalize_cover_letter_paragraphs(paragraphs)[:paragraph_cap]
    if not compacted:
        return []

    normalized = []
    for index, paragraph in enumerate(compacted):
        is_closing_paragraph = index == len(compacted) - 1 and any(
            token in paragraph.lower()
            for token in ("thank", "appreciate", "welcome the opportunity", "look forward")
        )
        normalized.append(
            _compact_cover_letter_paragraph(
                paragraph,
                max_sentences=closing_max_sentences if is_closing_paragraph else body_max_sentences,
                max_chars=closing_max_chars if is_closing_paragraph else body_max_chars,
            )
        )

    total_chars = sum(len(paragraph) for paragraph in normalized)
    while total_chars > total_target_chars:
        longest_index = max(range(len(normalized)), key=lambda idx: len(normalized[idx]))
        current = normalized[longest_index]
        tighter_limit = max(160, len(current) - 70)
        updated = _compact_cover_letter_paragraph(
            current,
            max_sentences=max(2, body_max_sentences - 1),
            max_chars=tighter_limit,
        )
        if updated == current:
            break
        normalized[longest_index] = updated
        total_chars = sum(len(paragraph) for paragraph in normalized)

    return [paragraph for paragraph in normalized if paragraph]


def _format_resume_skills(skills) -> str:
    if isinstance(skills, str):
        return skills
    if isinstance(skills, list):
        return _list_to_text(skills)
    if isinstance(skills, dict):
        groups = []
        for label, value in skills.items():
            value_text = _list_to_text(value)
            if value_text:
                groups.append(f"{label}: {value_text}")
        return "; ".join(groups)
    return _list_to_text(skills)


def _format_education_entries(education) -> str:
    if isinstance(education, str):
        return education
    if isinstance(education, dict):
        education = [education]
    if isinstance(education, list):
        entries = []
        for item in education:
            if isinstance(item, dict):
                main_parts = [item.get("degree"), item.get("institution")]
                main_text = ", ".join(str(part).strip() for part in main_parts if part)
                detail_parts = [
                    item.get("location"),
                    item.get("dates"),
                    f"GPA {item['gpa']}" if item.get("gpa") else None,
                    f"Average {item['average']}" if item.get("average") else None,
                ]
                details = " | ".join(str(part).strip() for part in detail_parts if part)
                awards_text = _list_to_text(item.get("awards"))
                if awards_text:
                    details = f"{details} | Awards: {awards_text}" if details else f"Awards: {awards_text}"
                if details:
                    entries.append(f"{main_text} ({details})" if main_text else details)
                elif main_text:
                    entries.append(main_text)
            else:
                item_text = str(item).strip()
                if item_text:
                    entries.append(item_text)
        return "; ".join(entries)
    return _list_to_text(education)


def _normalize_resume_experiences(experiences) -> list[dict]:
    if not isinstance(experiences, list):
        return []

    normalized = []
    for item in experiences:
        if not isinstance(item, dict):
            continue

        bullets = item.get("bullets", [])
        if isinstance(bullets, str):
            bullet_list = [bullets]
        elif isinstance(bullets, list):
            bullet_list = [str(bullet).strip() for bullet in bullets if str(bullet).strip()]
        else:
            bullet_text = str(bullets).strip()
            bullet_list = [bullet_text] if bullet_text else []

        normalized.append(
            {
                "company": str(item.get("company", "")).strip(),
                "role": str(item.get("role", "")).strip(),
                "dates": str(item.get("dates", "")).strip(),
                "bullets": bullet_list,
            }
        )

    return normalized


def normalize_document_sections(doc_type: str, sections: dict) -> dict:
    normalized = dict(sections or {})

    if doc_type == "resume":
        normalized["name"] = _list_to_text(normalized.get("name"))
        normalized["title"] = _list_to_text(normalized.get("title"))
        normalized["summary"] = _list_to_text(normalized.get("summary"))
        normalized["skills"] = _format_resume_skills(normalized.get("skills"))
        normalized["education"] = _format_education_entries(normalized.get("education"))
        normalized["experiences"] = _normalize_resume_experiences(normalized.get("experiences"))
        normalized["theme_id"] = _list_to_text(normalized.get("theme_id") or normalized.get("theme"))
        normalized["layout_strategy"] = _list_to_text(normalized.get("layout_strategy"))
        return normalized

    if doc_type == "cover_letter":
        normalized["name"] = _list_to_text(normalized.get("name"))
        normalized["date"] = datetime.now().strftime("%B %d, %Y").replace(" 0", " ")
        normalized["hiring_manager"] = _list_to_text(normalized.get("hiring_manager")) or "Hiring Manager"
        normalized["company"] = _list_to_text(normalized.get("company"))
        normalized["role"] = _list_to_text(normalized.get("role"))
        normalized["paragraphs"] = _normalize_cover_letter_paragraphs(normalized.get("paragraphs"))
        normalized["theme_id"] = _list_to_text(normalized.get("theme_id") or normalized.get("theme"))
        normalized["layout_strategy"] = _list_to_text(normalized.get("layout_strategy"))
        return normalized

    return normalized


def _choose_theme(doc_type: str, normalized_sections: dict) -> ThemeSpec:
    requested = normalized_sections.get("theme_id")
    if requested in THEMES:
        return THEMES[requested]

    requested_strategy = _requested_layout_strategy(normalized_sections)
    has_leadership_signal = _has_leadership_signal(normalized_sections)
    has_ats_signal = _has_ats_simplicity_signal(normalized_sections)
    has_design_signal = _has_design_signal(normalized_sections)

    if doc_type == "cover_letter":
        total_chars = sum(len(paragraph) for paragraph in normalized_sections.get("paragraphs", []))
        if requested_strategy == "compact":
            return THEMES["technical_compact"]
        if requested_strategy == "executive":
            return THEMES["executive_clean"]
        if requested_strategy == "balanced":
            return THEMES["classic_professional"]
        if requested_strategy == "ats_safe":
            return THEMES["technical_compact"] if total_chars > 780 else THEMES["ats_minimal"]
        if requested_strategy == "creative_safe":
            return THEMES["technical_compact"] if total_chars > 780 else THEMES["modern_minimal"]
        if total_chars > 780:
            return THEMES["technical_compact"]
        if has_ats_signal:
            return THEMES["ats_minimal"]
        if has_leadership_signal:
            return THEMES["executive_clean"]
        if has_design_signal:
            return THEMES["modern_minimal"]
        return THEMES["classic_professional"]

    experiences = normalized_sections.get("experiences", [])
    bullet_count = sum(len(item.get("bullets", [])) for item in experiences if isinstance(item, dict))
    density_score = (
        len(normalized_sections.get("summary", "")) // 80
        + len(normalized_sections.get("skills", "")) // 70
        + len(experiences) * 2
        + bullet_count
    )
    if requested_strategy == "compact":
        return THEMES["technical_compact"]
    if requested_strategy == "executive":
        return THEMES["executive_clean"]
    if requested_strategy == "balanced":
        return THEMES["classic_professional"]
    if requested_strategy == "ats_safe":
        if density_score >= 12:
            return THEMES["technical_compact"]
        return THEMES["ats_minimal"]
    if requested_strategy == "creative_safe":
        if density_score >= 12:
            return THEMES["technical_compact"]
        return THEMES["modern_minimal"]
    if density_score >= 12:
        return THEMES["technical_compact"]
    if has_ats_signal:
        return THEMES["ats_minimal"]
    if has_leadership_signal:
        return THEMES["executive_clean"]
    if has_design_signal:
        return THEMES["modern_minimal"]
    if density_score >= 10:
        return THEMES["technical_compact"]
    return THEMES["classic_professional"]


def _plan_resume_sections(normalized_sections: dict, theme: ThemeSpec) -> dict:
    return _plan_resume_sections_with_repairs(normalized_sections, theme, repair_level=0)


def _resume_summary_target(theme: ThemeSpec, repair_level: int) -> int:
    return max(140, theme.summary_target_chars - (repair_level * 40))


def _resume_skills_target(theme: ThemeSpec, repair_level: int) -> int:
    return max(150, theme.skills_target_chars - (repair_level * 35))


def _resume_max_experiences(theme: ThemeSpec, repair_level: int) -> int:
    reduction = 1 if repair_level >= 3 else 0
    return max(2, theme.max_resume_experiences - reduction)


def _resume_max_bullets(theme: ThemeSpec, repair_level: int) -> int:
    if repair_level >= 2:
        return 1
    return theme.max_bullets_per_experience


def _plan_resume_sections_with_repairs(
    normalized_sections: dict,
    theme: ThemeSpec,
    *,
    repair_level: int,
) -> dict:
    experiences = []
    for item in normalized_sections.get("experiences", [])[: _resume_max_experiences(theme, repair_level)]:
        bullets = item.get("bullets", [])[: _resume_max_bullets(theme, repair_level)]
        experiences.append(
            {
                "company": item.get("company", ""),
                "role": item.get("role", ""),
                "dates": item.get("dates", ""),
                "bullets": [_trim_text(bullet, max(85, 120 - (repair_level * 10))) for bullet in bullets],
            }
        )

    return {
        **normalized_sections,
        "summary": _trim_text(
            normalized_sections.get("summary", ""),
            _resume_summary_target(theme, repair_level),
        ),
        "skills": _trim_text(
            normalized_sections.get("skills", ""),
            _resume_skills_target(theme, repair_level),
        ),
        "experiences": experiences,
    }


def _plan_cover_letter_sections(normalized_sections: dict, theme: ThemeSpec) -> dict:
    return _plan_cover_letter_sections_with_repairs(normalized_sections, theme, repair_level=0)


def _plan_cover_letter_sections_with_repairs(
    normalized_sections: dict,
    theme: ThemeSpec,
    *,
    repair_level: int,
) -> dict:
    paragraph_cap = 4 if repair_level < 2 else 3
    body_max_sentences = 3 if repair_level < 2 else 2
    total_target_chars = max(620, COVER_LETTER_TARGET_TOTAL_CHARS - (repair_level * 150))
    body_max_chars = max(230, (320 if theme.density == "compact" else 360) - (repair_level * 35))
    closing_max_chars = max(150, (200 if theme.density == "compact" else 220) - (repair_level * 20))
    paragraphs = _fit_cover_letter_paragraphs(
        normalized_sections.get("paragraphs"),
        paragraph_cap=paragraph_cap,
        total_target_chars=total_target_chars,
        body_max_sentences=body_max_sentences,
        closing_max_sentences=2,
        body_max_chars=body_max_chars,
        closing_max_chars=closing_max_chars,
    )

    return {
        **normalized_sections,
        "paragraphs": paragraphs,
    }


def _estimate_resume_page_load(sections: dict) -> float:
    experiences = sections.get("experiences", [])
    bullet_count = sum(len(item.get("bullets", [])) for item in experiences)
    return (
        18
        + (len(sections.get("summary", "")) / 18)
        + (len(sections.get("skills", "")) / 18)
        + (len(sections.get("education", "")) / 32)
        + (len(experiences) * 8)
        + (bullet_count * 4)
    )


def _estimate_cover_letter_page_load(sections: dict) -> float:
    body_chars = sum(len(paragraph) for paragraph in sections.get("paragraphs", []))
    return (
        22
        + (body_chars / 20)
        + (len(sections.get("paragraphs", [])) * 4)
        + (len(sections.get("company", "")) / 18)
        + (len(sections.get("role", "")) / 18)
    )


def _verify_planned_document(
    doc_type: str,
    planned_sections: dict,
    theme: ThemeSpec,
    page_budget: int,
    layout_metrics: dict[str, Any],
) -> dict[str, Any]:
    issues: list[VerificationIssue] = []

    if doc_type == "resume":
        allowed_page_load = 92.0 if theme.density == "balanced" else 100.0
        estimated_page_load = _estimate_resume_page_load(planned_sections)
    else:
        allowed_page_load = 70.0 if theme.density == "balanced" else 76.0
        estimated_page_load = _estimate_cover_letter_page_load(planned_sections)

    if estimated_page_load > allowed_page_load:
        issues.append(
            VerificationIssue(
                code="page_budget_exceeded",
                message="Estimated layout exceeds the current page budget.",
                actual=round(estimated_page_load, 2),
                limit=allowed_page_load,
            )
        )

    if doc_type == "cover_letter" and len(planned_sections.get("paragraphs", [])) > 4:
        issues.append(
            VerificationIssue(
                code="paragraph_budget_exceeded",
                message="Cover letter body exceeds the supported paragraph budget.",
                actual=len(planned_sections.get("paragraphs", [])),
                limit=4,
            )
        )

    checks = {
        "page_budget": page_budget,
        "estimated_page_load": round(estimated_page_load, 2),
        "allowed_page_load": allowed_page_load,
        "page_load_ratio": round(estimated_page_load / allowed_page_load, 3),
        "within_page_budget": estimated_page_load <= allowed_page_load,
        "section_order_valid": True,
    }
    if doc_type == "cover_letter":
        checks["paragraph_count"] = len(planned_sections.get("paragraphs", []))

    return {
        "status": "passed" if not issues else "failed",
        "issues": [asdict(issue) for issue in issues],
        "checks": checks,
    }


def _next_repair_action(
    doc_type: str,
    theme: ThemeSpec,
    repair_level: int,
) -> tuple[RepairAction | None, str | None, int]:
    if theme.density != "compact":
        return (
            RepairAction(
                action="switch_theme",
                reason="Estimated page load exceeds the balanced theme budget.",
                details={"from_theme_id": theme.theme_id, "to_theme_id": "technical_compact"},
            ),
            "technical_compact",
            repair_level,
        )

    if doc_type == "resume":
        if repair_level == 0:
            return (
                RepairAction(
                    action="tighten_text_budgets",
                    reason="Estimated resume load still exceeds the one-page budget.",
                    details={
                        "summary_target_chars": _resume_summary_target(theme, 1),
                        "skills_target_chars": _resume_skills_target(theme, 1),
                    },
                ),
                None,
                1,
            )
        if repair_level == 1:
            return (
                RepairAction(
                    action="reduce_bullets",
                    reason="Estimated resume load still exceeds the one-page budget.",
                    details={"max_bullets_per_experience": _resume_max_bullets(theme, 2)},
                ),
                None,
                2,
            )
        if repair_level == 2:
            return (
                RepairAction(
                    action="drop_low_priority_experience",
                    reason="Estimated resume load still exceeds the one-page budget.",
                    details={"max_resume_experiences": _resume_max_experiences(theme, 3)},
                ),
                None,
                3,
            )
        return None, None, repair_level

    if repair_level == 0:
        return (
            RepairAction(
                action="tighten_cover_letter_budgets",
                reason="Estimated cover letter load still exceeds the one-page budget.",
                details={"body_paragraph_max_chars": max(230, 320 - 35)},
            ),
            None,
            1,
        )
    if repair_level == 1:
        return (
            RepairAction(
                action="reduce_cover_letter_paragraphs",
                reason="Estimated cover letter load still exceeds the one-page budget.",
                details={"max_paragraphs": 3},
            ),
            None,
            2,
        )
    if repair_level == 2:
        return (
            RepairAction(
                action="tighten_cover_letter_text",
                reason="Estimated cover letter load still exceeds the one-page budget.",
                details={"body_paragraph_max_chars": max(230, 320 - 105)},
            ),
            None,
            3,
        )
    return None, None, repair_level


def _build_document_plan_once(
    doc_type: str,
    normalized_sections: dict,
    theme: ThemeSpec,
    *,
    repair_level: int,
    repair_history: list[RepairAction],
    attempt_count: int,
) -> DocumentPlan:

    if doc_type == "resume":
        planned_sections = _plan_resume_sections_with_repairs(
            normalized_sections,
            theme,
            repair_level=repair_level,
        )
        layout_metrics = {
            "experience_count": len(planned_sections.get("experiences", [])),
            "bullet_count": sum(len(item.get("bullets", [])) for item in planned_sections.get("experiences", [])),
            "summary_chars": len(planned_sections.get("summary", "")),
            "skills_chars": len(planned_sections.get("skills", "")),
            "repair_level": repair_level,
        }
        section_order = ["summary", "experience", "skills", "education"]
        page_budget = 1
    else:
        planned_sections = _plan_cover_letter_sections_with_repairs(
            normalized_sections,
            theme,
            repair_level=repair_level,
        )
        layout_metrics = {
            "paragraph_count": len(planned_sections.get("paragraphs", [])),
            "body_chars": sum(len(paragraph) for paragraph in planned_sections.get("paragraphs", [])),
            "repair_level": repair_level,
        }
        section_order = ["date", "recipient", "subject", "body", "signature"]
        page_budget = 1

    verification = _verify_planned_document(
        doc_type,
        planned_sections,
        theme,
        page_budget,
        layout_metrics,
    )

    return DocumentPlan(
        doc_type=doc_type,
        page_budget=page_budget,
        theme_id=theme.theme_id,
        density=theme.density,
        normalized_sections=planned_sections,
        section_order=section_order,
        layout_metrics=layout_metrics,
        verification=verification,
        repair_history=[asdict(action) for action in repair_history],
        attempt_count=attempt_count,
    )


def build_document_plan(doc_type: str, sections: dict) -> DocumentPlan:
    normalized_sections = normalize_document_sections(doc_type, sections)
    initial_theme = _choose_theme(doc_type, normalized_sections)
    theme_override = initial_theme.theme_id
    repair_level = 0
    repair_history: list[RepairAction] = []

    for attempt_count in range(1, MAX_PLAN_REPAIR_ATTEMPTS + 1):
        theme = THEMES[theme_override]
        plan = _build_document_plan_once(
            doc_type,
            normalized_sections,
            theme,
            repair_level=repair_level,
            repair_history=repair_history,
            attempt_count=attempt_count,
        )
        if plan.verification["status"] == "passed":
            return plan

        action, next_theme_override, next_repair_level = _next_repair_action(
            doc_type,
            theme,
            repair_level,
        )
        if action is None:
            return plan

        repair_history.append(action)
        if next_theme_override is not None:
            theme_override = next_theme_override
        repair_level = next_repair_level

    return plan


def serialize_document_plan(plan: DocumentPlan) -> dict[str, Any]:
    return asdict(plan)


def _apply_theme(document: Document, theme: ThemeSpec):
    section = document.sections[0]
    section.top_margin = Inches(theme.top_margin_in)
    section.right_margin = Inches(theme.right_margin_in)
    section.bottom_margin = Inches(theme.bottom_margin_in)
    section.left_margin = Inches(theme.left_margin_in)

    style = document.styles["Normal"]
    style.font.name = theme.body_font
    style.font.size = Pt(theme.body_size_pt)
    style.font.color.rgb = RGBColor(*theme.body_color)
    style.paragraph_format.line_spacing = theme.line_spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(theme.paragraph_after_pt)


def _set_paragraph_bottom_border(paragraph, color: tuple[int, int, int], size: int = 8, space: int = 1):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), "".join(f"{component:02X}" for component in color))


def _add_header_divider(document: Document, theme: ThemeSpec):
    if not theme.header_divider:
        return
    divider = document.add_paragraph()
    divider.paragraph_format.space_before = Pt(1)
    divider.paragraph_format.space_after = Pt(theme.section_before_pt - 2)
    _set_paragraph_bottom_border(divider, theme.divider_color)


def _add_section_heading(document: Document, text: str, theme: ThemeSpec):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(theme.section_before_pt)
    paragraph.paragraph_format.space_after = Pt(theme.section_after_pt)
    paragraph.paragraph_format.keep_with_next = True
    heading_text = text.upper() if theme.heading_case == "upper" else text
    run = paragraph.add_run(heading_text)
    run.bold = True
    run.font.name = theme.heading_font
    run.font.size = Pt(theme.heading_size_pt)
    run.font.color.rgb = RGBColor(*theme.heading_color)


def _add_body_paragraph(document: Document, text: str, theme: ThemeSpec, *, indent_in: float = 0.0):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(theme.paragraph_after_pt)
    paragraph.paragraph_format.line_spacing = theme.line_spacing
    if indent_in:
        paragraph.paragraph_format.left_indent = Inches(indent_in)
    run = paragraph.add_run(text)
    run.font.name = theme.body_font
    run.font.size = Pt(theme.body_size_pt)
    run.font.color.rgb = RGBColor(*theme.body_color)
    return paragraph


def _render_resume(plan: DocumentPlan) -> bytes:
    theme = THEMES[plan.theme_id]
    sections = plan.normalized_sections

    document = Document()
    _apply_theme(document, theme)

    header = document.add_paragraph()
    header.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if theme.resume_header_alignment == "center"
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    header.paragraph_format.space_after = Pt(4)
    header.paragraph_format.keep_with_next = True
    name_run = header.add_run(sections.get("name", ""))
    name_run.bold = True
    name_run.font.name = theme.heading_font
    name_run.font.size = Pt(theme.name_size_pt)
    name_run.font.color.rgb = RGBColor(*theme.heading_color)

    title = document.add_paragraph()
    title.alignment = header.alignment
    title.paragraph_format.space_after = Pt(theme.section_before_pt)
    title_run = title.add_run(sections.get("title", ""))
    title_run.font.name = theme.body_font
    title_run.font.size = Pt(theme.title_size_pt)
    title_run.font.color.rgb = RGBColor(*theme.title_color)
    title_run.italic = theme.title_italic
    _add_header_divider(document, theme)

    if sections.get("summary"):
        _add_section_heading(document, "Summary", theme)
        _add_body_paragraph(document, sections["summary"], theme)

    experiences = sections.get("experiences", [])
    if experiences:
        _add_section_heading(document, "Experience", theme)
        for experience in experiences:
            line = document.add_paragraph()
            line.paragraph_format.keep_with_next = True
            line.paragraph_format.space_before = Pt(0)
            line.paragraph_format.space_after = Pt(theme.paragraph_after_pt)
            line.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)

            role_run = line.add_run(experience.get("role", ""))
            role_run.bold = True
            role_run.font.size = Pt(theme.body_size_pt)
            role_run.font.name = theme.body_font
            company = experience.get("company", "")
            if company:
                company_run = line.add_run(f" | {company}")
                company_run.font.size = Pt(theme.body_size_pt)
                company_run.font.name = theme.body_font
            dates = experience.get("dates", "")
            if dates:
                dates_run = line.add_run(f"\t{dates}")
                dates_run.font.size = Pt(theme.body_size_pt)
                dates_run.font.name = theme.body_font

            for bullet in experience.get("bullets", []):
                _add_body_paragraph(document, f"• {bullet}", theme, indent_in=theme.bullet_indent_in)

    if sections.get("skills"):
        _add_section_heading(document, "Skills", theme)
        _add_body_paragraph(document, sections["skills"], theme)

    if sections.get("education"):
        _add_section_heading(document, "Education", theme)
        _add_body_paragraph(document, sections["education"], theme)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_cover_letter(plan: DocumentPlan) -> bytes:
    theme = THEMES[plan.theme_id]
    sections = plan.normalized_sections

    document = Document()
    _apply_theme(document, theme)

    date_paragraph = _add_body_paragraph(document, sections.get("date", ""), theme)
    date_paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
        if theme.cover_letter_date_alignment == "right"
        else WD_ALIGN_PARAGRAPH.LEFT
    )

    recipient_lines = [line for line in [sections.get("hiring_manager", ""), sections.get("company", "")] if line]
    for index, line_text in enumerate(recipient_lines):
        recipient = document.add_paragraph()
        recipient.paragraph_format.space_before = Pt(theme.section_before_pt if index == 0 else 0)
        recipient.paragraph_format.space_after = Pt(theme.paragraph_after_pt if index == len(recipient_lines) - 1 else 0)
        recipient.paragraph_format.keep_with_next = True
        line = recipient.add_run(line_text)
        line.font.name = theme.body_font
        line.font.size = Pt(theme.body_size_pt)

    subject = document.add_paragraph()
    subject.paragraph_format.space_before = Pt(theme.section_before_pt)
    subject.paragraph_format.space_after = Pt(theme.section_before_pt)
    subject.paragraph_format.keep_with_next = True
    subject_run = subject.add_run(f"Re: {sections.get('role', '')}")
    subject_run.bold = True
    subject_run.font.name = theme.body_font
    subject_run.font.size = Pt(theme.body_size_pt)
    if theme.header_divider:
        _add_header_divider(document, theme)

    for paragraph_text in sections.get("paragraphs", []):
        _add_body_paragraph(document, paragraph_text, theme)

    signature = document.add_paragraph()
    signature.paragraph_format.space_before = Pt(theme.section_before_pt)
    signature.paragraph_format.space_after = Pt(0)
    signature.paragraph_format.keep_with_next = True
    sig_run = signature.add_run("Sincerely,")
    sig_run.font.name = theme.body_font
    sig_run.font.size = Pt(theme.body_size_pt)

    name = document.add_paragraph()
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(0)
    name_run = name.add_run(sections.get("name", ""))
    name_run.font.name = theme.body_font
    name_run.font.size = Pt(theme.body_size_pt)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_document(plan: DocumentPlan) -> bytes:
    if plan.doc_type == "resume":
        return _render_resume(plan)
    if plan.doc_type == "cover_letter":
        return _render_cover_letter(plan)
    raise ValueError(f"Unsupported doc_type: {plan.doc_type}")
