import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import document_engine


def test_build_document_plan_respects_requested_theme():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Avery Carter",
            "title": "Backend Engineer",
            "summary": "Experienced backend engineer.",
            "skills": ["Python", "FastAPI", "PostgreSQL"],
            "education": "BSc Computer Science",
            "experiences": [],
            "theme_id": "technical_compact",
        },
    )

    assert plan.theme_id == "technical_compact"
    assert plan.density == "compact"
    assert plan.page_budget == 1
    assert plan.verification["status"] == "passed"
    assert plan.repair_history == []
    assert plan.attempt_count == 1


def test_build_document_plan_chooses_compact_theme_for_dense_resume():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Avery Carter",
            "title": "Backend Engineer",
            "summary": " ".join(["Built APIs and distributed systems."] * 12),
            "skills": {
                "Languages": ["Python", "TypeScript", "SQL", "JavaScript"],
                "Cloud": ["AWS", "GCP", "Azure"],
                "Data": ["PostgreSQL", "MongoDB", "Redis"],
            },
            "education": "BSc Computer Science",
            "experiences": [
                {"company": "A", "role": "One", "dates": "2020-2021", "bullets": ["A", "B", "C"]},
                {"company": "B", "role": "Two", "dates": "2021-2022", "bullets": ["A", "B", "C"]},
                {"company": "C", "role": "Three", "dates": "2022-2023", "bullets": ["A", "B", "C"]},
                {"company": "D", "role": "Four", "dates": "2023-2024", "bullets": ["A", "B", "C"]},
            ],
        },
    )

    assert plan.theme_id == "technical_compact"
    assert plan.verification["status"] == "passed"
    assert plan.layout_metrics["experience_count"] <= 4
    assert plan.layout_metrics["bullet_count"] <= 8


def test_build_document_plan_chooses_executive_theme_for_leadership_resume():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Morgan Lee",
            "title": "Staff Platform Engineer",
            "summary": (
                "Staff platform engineer with experience leading reliability programs, "
                "mentoring engineers, and shaping backend platform strategy."
            ),
            "skills": {
                "Languages": ["Python", "Go", "SQL"],
                "Cloud": ["AWS", "GCP"],
            },
            "education": "BSc Computer Science",
            "experiences": [
                {
                    "company": "Northstar Systems",
                    "role": "Staff Platform Engineer",
                    "dates": "2022 to Present",
                    "bullets": [
                        "Led platform reliability initiatives across core services.",
                        "Mentored backend engineers and drove architecture reviews.",
                    ],
                },
                {
                    "company": "Atlas Data",
                    "role": "Senior Backend Engineer",
                    "dates": "2019 to 2022",
                    "bullets": [
                        "Built shared APIs and internal tooling for multiple teams.",
                        "Improved on-call stability and deployment safety.",
                    ],
                },
            ],
        },
    )

    assert plan.theme_id == "executive_clean"
    assert plan.verification["status"] == "passed"
    assert plan.repair_history == []


def test_build_document_plan_chooses_ats_minimal_for_ats_safe_strategy():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Taylor Brooks",
            "title": "Operations Coordinator",
            "layout_strategy": "ats_safe",
            "summary": "Operations coordinator with experience supporting scheduling, compliance, and cross-team administration workflows.",
            "skills": {
                "Tools": ["Excel", "Google Workspace", "CRM"],
                "Strengths": ["Scheduling", "Documentation", "Reporting"],
            },
            "education": "Diploma in Business Administration",
            "experiences": [
                {
                    "company": "Northview Health",
                    "role": "Operations Coordinator",
                    "dates": "2022 to Present",
                    "bullets": [
                        "Coordinated scheduling and documentation across multiple teams.",
                        "Maintained compliance records and reporting workflows.",
                    ],
                }
            ],
        },
    )

    assert plan.theme_id == "ats_minimal"
    assert plan.verification["status"] == "passed"
    assert plan.repair_history == []


def test_build_document_plan_chooses_modern_minimal_for_design_role():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Jordan Vale",
            "title": "Product Designer",
            "summary": (
                "Product designer focused on UX systems, interaction design, and "
                "clear product storytelling across web and mobile surfaces."
            ),
            "skills": {
                "Design": ["Figma", "Prototyping", "Design systems"],
                "Product": ["UX research", "Interaction design", "Accessibility"],
            },
            "education": "BDes Interaction Design",
            "experiences": [
                {
                    "company": "North Coast",
                    "role": "Senior Product Designer",
                    "dates": "2022 to Present",
                    "bullets": [
                        "Led UX design for core product flows and design system work.",
                        "Partnered with product and engineering on shipped interaction details.",
                    ],
                }
            ],
        },
    )

    assert plan.theme_id == "modern_minimal"
    assert plan.verification["status"] == "passed"
    assert plan.repair_history == []


def test_build_document_plan_repairs_dense_resume_until_within_budget():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Avery Carter",
            "title": "Senior Backend Engineer",
            "theme_id": "classic_professional",
            "summary": " ".join(["Built high-volume APIs, ingestion systems, and internal developer platforms."] * 12),
            "skills": {
                "Languages": ["Python", "TypeScript", "SQL", "Go", "JavaScript"],
                "Cloud": ["AWS", "GCP", "Azure", "Cloud Run", "Kubernetes"],
                "Data": ["PostgreSQL", "Redis", "MongoDB", "BigQuery"],
                "Other": ["FastAPI", "Flask", "CI/CD", "Observability", "Terraform"],
            },
            "education": "BSc Computer Science, Memorial University (2021-2025 | GPA 3.83)",
            "experiences": [
                {"company": "Nexa Labs", "role": "Senior Backend Engineer", "dates": "2023 to Present", "bullets": ["Built resilient APIs for document generation and matching systems.", "Led backend reliability work across storage, auth, and orchestration.", "Improved latency and observability."]},
                {"company": "Orbit Systems", "role": "Platform Engineer", "dates": "2021 to 2023", "bullets": ["Built Python services for internal tools.", "Improved CI/CD and deployment automation.", "Partnered with product teams."]},
                {"company": "Northwind", "role": "Software Engineer", "dates": "2020 to 2021", "bullets": ["Implemented APIs and data pipelines.", "Improved database design.", "Supported production operations."]},
                {"company": "Acme", "role": "Engineer", "dates": "2019 to 2020", "bullets": ["Built backend features.", "Maintained integrations.", "Wrote tests and docs."]},
            ],
        },
    )

    assert plan.theme_id == "technical_compact"
    assert plan.verification["status"] == "passed"
    assert plan.attempt_count > 1
    assert [item["action"] for item in plan.repair_history]
    assert plan.layout_metrics["bullet_count"] <= 4
    assert plan.layout_metrics["experience_count"] <= 4


def test_build_document_plan_compacts_cover_letter():
    plan = document_engine.build_document_plan(
        "cover_letter",
        {
            "name": "Avery Carter",
            "company": "Boam AI",
            "role": "Backend Engineer",
            "paragraphs": [
                "Sentence one. Sentence two. Sentence three. Sentence four.",
                "This paragraph is intentionally very long so that the engine must trim it while still keeping the cover letter useful, direct, and aligned with the one-page target that the planner is enforcing for the final render.",
                "Third paragraph stays in place.",
                "Thank you for your time. I appreciate your consideration. I look forward to speaking with you soon.",
                "This one should be dropped completely.",
            ],
        },
    )

    assert plan.page_budget == 1
    assert plan.layout_metrics["paragraph_count"] == 4
    assert plan.normalized_sections["paragraphs"][0] == "Sentence one. Sentence two. Sentence three."
    assert plan.normalized_sections["paragraphs"][-1] == "Thank you for your time. I appreciate your consideration."


def test_build_document_plan_repairs_dense_cover_letter_until_within_budget():
    plan = document_engine.build_document_plan(
        "cover_letter",
        {
            "name": "Avery Carter",
            "company": "Boam AI",
            "role": "Backend Engineer",
            "theme_id": "classic_professional",
            "paragraphs": [
                " ".join(["I build backend systems that support job matching, document generation, and production workflows."] * 4),
                " ".join(["At Nexa Labs, I led API platform work, database integrations, and reliability improvements across multiple services."] * 4),
                " ".join(["I would bring strong Python, FastAPI, PostgreSQL, and infrastructure experience to this role while staying close to product delivery needs."] * 4),
                " ".join(["Thank you for your time and consideration. I would welcome the opportunity to discuss how I could contribute."] * 3),
            ],
        },
    )

    assert plan.verification["status"] == "passed"
    assert plan.attempt_count > 1
    assert [item["action"] for item in plan.repair_history]
    assert len(plan.normalized_sections["paragraphs"]) <= 3


def test_render_document_outputs_expected_cover_letter_structure():
    plan = document_engine.build_document_plan(
        "cover_letter",
        {
            "name": "Avery Carter",
            "company": "Boam AI",
            "role": "Backend Engineer",
            "paragraphs": [
                "I am excited to apply.",
                "I have built backend systems with FastAPI and PostgreSQL.",
                "I would bring strong API and distributed-systems experience to the role.",
                "Thank you for your time and consideration.",
            ],
        },
    )

    rendered = document_engine.render_document(plan)
    document = Document(io.BytesIO(rendered))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    assert paragraphs[0]
    assert paragraphs[1] == "Hiring Manager"
    assert paragraphs[2] == "Boam AI"
    assert paragraphs[3] == "Re: Backend Engineer"
    assert paragraphs[-2] == "Sincerely,"
    assert paragraphs[-1] == "Avery Carter"


def test_render_document_applies_executive_resume_header_alignment():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Morgan Lee",
            "title": "Staff Platform Engineer",
            "summary": "Led platform strategy, reliability, and mentoring initiatives.",
            "skills": "Python, Go, AWS",
            "education": "BSc Computer Science",
            "experiences": [
                {
                    "company": "Northstar Systems",
                    "role": "Staff Platform Engineer",
                    "dates": "2022 to Present",
                    "bullets": ["Led reliability programs.", "Mentored backend engineers."],
                }
            ],
        },
    )

    rendered = document_engine.render_document(plan)
    document = Document(io.BytesIO(rendered))

    assert plan.theme_id == "executive_clean"
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert document.paragraphs[2].text == "SUMMARY"


def test_render_document_applies_executive_cover_letter_date_alignment():
    plan = document_engine.build_document_plan(
        "cover_letter",
        {
            "name": "Morgan Lee",
            "company": "Northstar Systems",
            "role": "Staff Platform Engineer",
            "paragraphs": [
                "I am excited to apply for the Staff Platform Engineer role at Northstar Systems.",
                "I have led platform reliability programs and backend strategy across distributed systems teams.",
                "I would bring strong technical leadership and mentoring experience to the role.",
                "Thank you for your time and consideration.",
            ],
        },
    )

    rendered = document_engine.render_document(plan)
    document = Document(io.BytesIO(rendered))

    assert plan.theme_id == "executive_clean"
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_render_document_applies_ats_minimal_header_alignment():
    plan = document_engine.build_document_plan(
        "resume",
        {
            "name": "Taylor Brooks",
            "title": "Operations Coordinator",
            "layout_strategy": "ats_safe",
            "summary": "Supports operations, scheduling, and documentation workflows.",
            "skills": "Excel, Reporting, Scheduling",
            "education": "Diploma in Business Administration",
            "experiences": [
                {
                    "company": "Northview Health",
                    "role": "Operations Coordinator",
                    "dates": "2022 to Present",
                    "bullets": ["Coordinated schedules.", "Maintained compliance records."],
                }
            ],
        },
    )

    rendered = document_engine.render_document(plan)
    document = Document(io.BytesIO(rendered))

    assert plan.theme_id == "ats_minimal"
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert document.paragraphs[2].text == "Summary"


def test_render_document_applies_modern_minimal_cover_letter_date_alignment():
    plan = document_engine.build_document_plan(
        "cover_letter",
        {
            "name": "Jordan Vale",
            "company": "North Coast",
            "role": "Product Designer",
            "layout_strategy": "creative_safe",
            "paragraphs": [
                "I am excited to apply for the Product Designer role at North Coast.",
                "I design thoughtful UX flows, systems, and interaction details for product teams.",
                "I would bring strong product thinking, visual clarity, and collaboration to the role.",
                "Thank you for your time and consideration.",
            ],
        },
    )

    rendered = document_engine.render_document(plan)
    document = Document(io.BytesIO(rendered))

    assert plan.theme_id == "modern_minimal"
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
