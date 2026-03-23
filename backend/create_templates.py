"""Generate initial .docx templates with Jinja2 placeholders for docxtpl."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os


def create_resume_template():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Name header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ name }}")
    run.bold = True
    run.font.size = Pt(24)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ title }}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Summary
    doc.add_heading("Summary", level=2)
    doc.add_paragraph("{{ summary }}")

    # Experience
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("{% for exp in experiences %}")
    p = doc.add_paragraph()
    run = p.add_run("{{ exp.role }}")
    run.bold = True
    p.add_run(" | {{ exp.company }} | {{ exp.dates }}")
    doc.add_paragraph("{% for bullet in exp.bullets %}")
    doc.add_paragraph("• {{ bullet }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")

    # Skills
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("{{ skills }}")

    # Education
    doc.add_heading("Education", level=2)
    doc.add_paragraph("{{ education }}")

    os.makedirs("templates", exist_ok=True)
    doc.save("templates/resume.docx")
    print("Created templates/resume.docx")


def create_cover_letter_template():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Date and header
    doc.add_paragraph("{{ date }}")
    doc.add_paragraph()
    doc.add_paragraph("{{ hiring_manager }}")
    doc.add_paragraph("{{ company }}")
    doc.add_paragraph()
    doc.add_paragraph("Re: {{ role }}")
    doc.add_paragraph()

    # Body paragraphs
    doc.add_paragraph("{% for para in paragraphs %}")
    doc.add_paragraph("{{ para }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("{{ name }}")

    os.makedirs("templates", exist_ok=True)
    doc.save("templates/cover_letter.docx")
    print("Created templates/cover_letter.docx")


if __name__ == "__main__":
    create_resume_template()
    create_cover_letter_template()
